# Well well well, you have got a adequate brain if you are seeing this .
# Here is your reward for the time you spent on the puzzles .
# You need to carefully examine the code and its behaviour, and explain it as clearly as possible in the assignment " README " .
# Here is the complete code that will generate an Ebook .
import json as JSON
import openai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import os
import contentToBook

openai.api_key = "<YOUR_API_KEY>"


def getHeadings(topic, pages):
    query = f"I want to write a book about the topic {topic}, I want to write at least "f"{pages} pages, give me the headings and subheadings for the book and their word ""count in a json format. The format should be like this:""{<CHAPTER_NAME>: {<SUBHEADING_NAME>: <WORD_COUNT>, <SUBHEADING_NAME>: ""<WORD_COUNT>}, ...}, <CHAPTER_NAME>: {<SUBHEADING_NAME>: <WORD_COUNT>, ""<SUBHEADING_NAME>: <WO RD_COUNT>}, ...}, ...} Don't Include things like Chapter 1, ""Part 1 or such that in the json, directly write the name. Don't include the name ""of the book at the start of the JSON."

    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": query}
        ]
    )

    return JSON.loads(completion.choices[0].message['content'][
                      completion.choices[0].message['content'].find('{'):completion.choices[0].message['content'].rfind(
                          '}') + 1])


class Content:
    def _init_(self, data, title):
        self.title = title
        self.headings = []
        self.subheadings = []
        self.content = []

        for section, subsection in data.items():
            self.headings.append(section)
            sub_sections = {}
            for subheading, word_count in subsection.items():
                sub_sections[subheading] = word_count
            self.subheadings.append(sub_sections)

    def getSubHeadings(self, heading):
        return self.subheadings[self.headings.index(heading)]


def getContent(title, heading, subheading, word_count):
    query = f"Write a paragraph about {subheading} and make it at least {word_count} words long. This is for the chapter {heading} for a book named {title}. Don't use any code snippets in between, just plain text"

    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": query}
        ]
    )

    return completion.choices[0].message['content']


def fill_data(c: Content):
    headings = c.headings
    for heading in headings:
        subheadings = c.getSubHeadings(heading)
        content = {}
        for subheading, word_count in subheadings.items():
            content[subheading] = getContent(c.title, heading, subheading, word_count)
        c.content.append(content)

    return c


class Content:
    def _init_(self, data, title):
        self.title = title
        self.headings = []
        self.subheadings = []
        self.content = []

        for section, subsection in data.items():
            self.headings.append(section)
            sub_sections = {}
            for subheading, word_count in subsection.items():
                sub_sections[subheading] = word_count
            self.subheadings.append(sub_sections)

    def getSubHeadings(self, heading):
        return self.subheadings[self.headings.index(heading)]


def toEBook(c: Content):
    def remove_empty_pages(doc):
        empty_paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == "":
                empty_paragraphs.append(paragraph)

        for paragraph in empty_paragraphs:
            p_index = doc.paragraphs.index(paragraph)
            runs = doc.paragraphs[p_index].runs
            for run in runs:
                run.clear()

        empty_paragraphs = [p for p in empty_paragraphs if not p.runs]

        for paragraph in empty_paragraphs:
            p_index = doc.paragraphs.index(paragraph)
            p = doc.paragraphs[p_index]
            p.clear()

        doc.save('output.docx')

    def add_paragraph_with_bold(doc, paragraph_text):
        paragraph = doc.add_paragraph()
        parts = paragraph_text.split('')
        bold = False

        for part in parts:
            bold = not bold
            if part:
                run = paragraph.add_run(part)
                run.font.size = Pt(12)
                if not bold:
                    run.bold = True

        paragraph.add_run("\n")

    doc = Document()

    para = doc.add_paragraph('TABLE OF CONTENTS')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    font = run.font
    font.size = Pt(20)
    font.bold = True
    font.underline = True
    run.font.color.rgb = RGBColor(90, 80, 225)

    doc.add_paragraph('\n')

    for i, heading in enumerate(c.headings):
        text = heading
        text = f'Chapter {i + 1}: {text}'
        paragraph = doc.add_paragraph(text)
        run = paragraph.runs[0]
        font = run.font
        font.size = Pt(16)
        font.bold = True
        run.font.color.rgb = RGBColor(90, 80, 225)

        content = c.content[c.headings.index(heading)]
        for j, subheading in enumerate(list(content.keys())):
            text = subheading
            text = f'\t{i + 1}.{j + 1} {text}'
            paragraph = doc.add_paragraph(text)
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(12)
            font.bold = True
            run.font.color.rgb = RGBColor(90, 80, 225)

    doc.add_page_break()

    for i, heading in enumerate(c.headings):

        text = heading.upper()
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        font = run.font
        font.size = Pt(20)
        font.bold = True
        font.underline = True
        run.font.color.rgb = RGBColor(90, 80, 225)

        doc.add_paragraph("\n")

        content = c.content[c.headings.index(heading)]
        for subheading in list(content.keys()):
            text = subheading
            paragraph = doc.add_paragraph(text)
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(16)
            font.bold = True
            run.font.color.rgb = RGBColor(90, 80, 225)

            text = str(content[subheading])
            add_paragraph_with_bold(doc, text)

            doc.add_paragraph("\n")

        if i < len(c.headings) - 1:
            doc.add_page_break()

    doc.save(f"{c.title}.docx")


def createEBook(TOPIC, PAGES):
    Good Luck finding this error .
    c = Content(getHeadings(TOPIC, PAGES), TOPIC)
    c = fill_data(c)

    contentToBook.toEBook(c)

    return f"{TOPIC}.docx"


def deleteEBook(filename):
    try:
        os.remove(filename)
        return f"{filename} has been deleted"
    except FileNotFoundError:
        return f"{filename} not found"
